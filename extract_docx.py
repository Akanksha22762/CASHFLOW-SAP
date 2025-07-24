import docx
import sys

def extract_docx_text(filename):
    """Extract text from a Word document"""
    try:
        doc = docx.Document(filename)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also check for tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    except Exception as e:
        print(f"Error reading {filename}: {e}")
        return None

if __name__ == "__main__":
    filename = "Cash_Flow_Model_AINurturing parameters.docx"
    text = extract_docx_text(filename)
    
    if text:
        print("=== EXTRACTED CONTENT ===")
        print(text)
        
        # Save to text file
        output_file = "Cash_Flow_Model_AINurturing_parameters.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"\nContent saved to: {output_file}")
    else:
        print("Failed to extract content from the document.") 